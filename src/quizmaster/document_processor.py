import os
import io
from typing import Dict, List, Optional
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
import requests
import re
import logging
import json
import tiktoken
from contextgem import Document as ContextGemDocument, DocumentLLM, StringConcept, Aspect, NumericalConcept, DateConcept
import requests
from typing import Optional # Added for QuizConfig type hint
from .config import QuizConfig # Added for QuizConfig type hint

# Configure logging to suppress pdfminer warnings
pdfminer_logger = logging.getLogger('pdfminer')
pdfminer_logger.setLevel(logging.ERROR)
pdfpage_logger = logging.getLogger('pdfminer.pdfpage')
pdfpage_logger.setLevel(logging.ERROR)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, persist_directory: str = "documents", config: Optional[QuizConfig] = None):
        self.supported_formats = ['pdf', 'docx', 'txt', 'html']
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)
        self.config = config  # Store the passed config
        
        self.contextgem_llm: Optional[DocumentLLM] = None
        self._initialize_contextgem_llm() # Initialize ContextGem LLM
        
        # Simplified chunking settings
        self.max_chunk_size = 8000
        self.encoding = tiktoken.get_encoding("cl100k_base")

        # Define default concepts for extraction if ContextGem is used
        # Simplified for testing JSON output with smaller models
        self.default_concepts_for_cg = [
            StringConcept(name="Main Topic", description="Primary topic or subject of the document", singular_occurrence=True, add_references=False),
            StringConcept(name="Document Type", description="Type or category of the document (e.g., report, article, manual)", singular_occurrence=True, add_references=False),
            # StringConcept(name="Key Arguments", description="Main arguments or theses presented", singular_occurrence=False, add_references=False),
            # StringConcept(name="Conclusions", description="Main conclusions or outcomes reported", singular_occurrence=False, add_references=False),
            # StringConcept(name="Technical Terms", description="Specific technical terms or jargon used", singular_occurrence=False, add_references=False),
            # StringConcept(name="Key People", description="Important individuals mentioned", singular_occurrence=False, add_references=False),
            # StringConcept(name="Key Organizations", description="Important organizations or entities mentioned", singular_occurrence=False, add_references=False),
            # DateConcept(name="Key Dates", description="Significant dates or time periods mentioned", singular_occurrence=False, add_references=False),
        ]

    def _initialize_contextgem_llm(self):
        """Initialize ContextGem DocumentLLM instance, preferring UI selected model if OpenAI."""
        try:
            openai_api_key = os.environ.get("OPENAI_API_KEY")
            anthropic_api_key = os.environ.get("ANTHROPIC_API_KEY")

            # Priority 1: Use globally selected OpenAI model if available and an API key is present
            if self.config and self.config.current_model and self.config.current_model in self.config.openai_models:
                if openai_api_key:
                    cg_model_name = self.config.current_model
                    if not cg_model_name.startswith("openai/"):
                        # Ensure the model name is prefixed with 'openai/' for ContextGem
                        cg_model_name = f"openai/{cg_model_name}"
                    
                    self.contextgem_llm = DocumentLLM(model=cg_model_name, api_key=openai_api_key)
                    logger.info(f"DocumentProcessor using UI selected OpenAI model '{cg_model_name}' for ContextGem.")
                    return
                else:
                    logger.error(f"DocumentProcessor: UI selected OpenAI model {self.config.current_model} but OPENAI_API_KEY not found. Cannot use this model.")
                    # Do not proceed to Ollama if an OpenAI model was explicitly chosen but key is missing.
                    # Let it fall through to the end where self.contextgem_llm might remain None or be set by other fallbacks if any.
                    # For now, let's ensure it doesn't pick Ollama if OpenAI was the intent but failed due to key.
                    # The final fallbacks (hardcoded OpenAI/Anthropic) will also fail if keys are missing.
                    # The most robust here is to set to None and return, letting the calling code handle no LLM.
                    self.contextgem_llm = None
                    logger.warning("DocumentProcessor: ContextGem LLM set to None due to missing API key for selected OpenAI model.")
                    return # Explicitly stop further model search if user intended OpenAI but key is missing

            # Priority 2: Try Ollama
            ollama_base_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
            try:
                response = requests.get(f"{ollama_base_url}/api/tags", timeout=2)
                if response.status_code == 200:
                    models_data = response.json()
                    available_ollama_models = [model['name'] for model in models_data.get('models', [])]
                    logger.info(f"DocumentProcessor: Available local Ollama models: {available_ollama_models}")
                    
                    selected_ollama_model = None # Ensure it's reset
                    ui_selected_model = self.config.current_model if self.config else None
                    logger.info(f"DocumentProcessor: UI selected model (config.current_model): {ui_selected_model}")

                    # Check if UI selected model is an Ollama model and is available
                    if ui_selected_model and ui_selected_model not in self.config.openai_models and ui_selected_model in available_ollama_models:
                        selected_ollama_model = ui_selected_model
                        logger.info(f"DocumentProcessor: UI selected Ollama model '{selected_ollama_model}' is available locally.")
                    else:
                        if ui_selected_model and ui_selected_model not in self.config.openai_models:
                            logger.info(f"DocumentProcessor: UI selected Ollama model '{ui_selected_model}' is NOT available locally or not an Ollama model. Checking preferred list.")
                        # Fallback to preferred Ollama models
                        preferred_ollama_models = ["mistral:7b", "qwen2.5:7b", "llama3.3:8b", "deepseek-coder:6.7b"]
                        for model in preferred_ollama_models:
                            if model in available_ollama_models:
                                selected_ollama_model = model
                                logger.info(f"DocumentProcessor: Found preferred Ollama model '{selected_ollama_model}' locally.")
                                break
                        if not selected_ollama_model and available_ollama_models:
                            selected_ollama_model = available_ollama_models[0]
                            logger.info(f"DocumentProcessor: Using first available Ollama model '{selected_ollama_model}' locally.")
                    
                    if selected_ollama_model:
                        self.contextgem_llm = DocumentLLM(model=f"ollama/{selected_ollama_model}", api_base=ollama_base_url)
                        logger.info(f"DocumentProcessor using Ollama model '{selected_ollama_model}' for ContextGem.")
                        return
            except Exception as e:
                logger.info(f"DocumentProcessor: Ollama not available or error selecting model: {e}")

            # Priority 3: Fallback to hardcoded OpenAI model if key exists (and UI model wasn't OpenAI or key was missing for it)
            if openai_api_key:
                self.contextgem_llm = DocumentLLM(model="openai/gpt-4o-mini", api_key=openai_api_key)
                logger.info("DocumentProcessor using fallback OpenAI model 'gpt-4o-mini' for ContextGem.")
                return
            
            # Priority 4: Fallback to Anthropic
            if anthropic_api_key:
                self.contextgem_llm = DocumentLLM(model="anthropic/claude-3-5-sonnet", api_key=anthropic_api_key)
                logger.info("DocumentProcessor using fallback Anthropic model 'claude-3-5-sonnet' for ContextGem.")
                return

            logger.warning("DocumentProcessor: ContextGem LLM could not be initialized. No suitable model configuration found. Concept extraction might be limited.")
            self.contextgem_llm = None

        except Exception as e:
            logger.error(f"DocumentProcessor: Failed to initialize ContextGem LLM: {e}", exc_info=True)
            self.contextgem_llm = None

    def update_llm_configuration(self):
        """
        Re-initializes the ContextGem DocumentLLM instance based on the current
        configuration. This should be called if the global model config changes.
        """
        logger.info("DocumentProcessor attempting to update LLM configuration for ContextGem.")
        if not self.config:
            logger.warning("DocumentProcessor has no config set; LLM re-initialization might use defaults.")
        self._initialize_contextgem_llm()
    
    # get_current_model and set_model are removed as model selection for ContextGem
    # is handled during DocumentLLM initialization or should be managed by a central LLMManager if shared.
    # For now, DocumentProcessor uses its initialized ContextGem LLM.

    def count_tokens(self, text: str) -> int:
        """Count approximate tokens in text."""
        try:
            return len(self.encoding.encode(text))
        except Exception:
            # Fallback: rough estimation
            return len(text) // 4

    def smart_chunk_text(self, text: str) -> List[str]:
        """Efficiently chunk text for ContextGem processing."""
        token_count = self.count_tokens(text)
        
        if token_count <= self.max_chunk_size:
            return [text]
        
        # Simple paragraph-based chunking - let ContextGem handle the complexity
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            test_chunk = current_chunk + ("\n\n" if current_chunk else "") + paragraph
            
            if self.count_tokens(test_chunk) <= self.max_chunk_size:
                current_chunk = test_chunk
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text]

    def extract_concepts(self, text: str, filename: str, file_format: str) -> List[Dict]:
        """Extract concepts, trying ContextGem first, then a direct LLM fallback."""
        if not isinstance(text, str) or not text.strip():
            logger.warning("Empty or invalid text provided for concept extraction.")
            return []

        if not self.contextgem_llm:
            logger.warning("ContextGem LLM not initialized for primary extraction path. Attempting direct fallback.")
            # Try fallback directly if ContextGem LLM isn't even set up
            try:
                concepts = self._fallback_extract_basic(text, filename, file_format)
                if concepts:
                    return concepts
            except Exception as fallback_error:
                logger.error(f"Initial fallback extraction failed: {str(fallback_error)}", exc_info=True)
            return self._generate_default_concepts(filename, file_format) # Last resort

        # Try ContextGem first
        logger.info(f"Attempting concept extraction with ContextGem for {filename}...")
        contextgem_concepts = None
        try:
            contextgem_concepts = self._extract_concepts_contextgem(text, filename, file_format)
        except Exception as e: # Catch exceptions from within _extract_concepts_contextgem if any slip through its own try/except
            logger.error(f"Exception during _extract_concepts_contextgem call for {filename}: {str(e)}", exc_info=True)
            contextgem_concepts = None # Ensure it's None to trigger fallback

        if contextgem_concepts is not None and len(contextgem_concepts) > 0: # Check for None and also if it returned an empty list explicitly (meaning it ran but found nothing, which is different from failing)
            logger.info(f"Successfully extracted {len(contextgem_concepts)} concepts using ContextGem for {filename}.")
            return contextgem_concepts
        
        # If ContextGem returned None (signaling failure/no usable concepts) or an empty list (if we decide empty list also means try fallback)
        logger.warning(f"ContextGem did not yield usable concepts for {filename}. Attempting direct LLM fallback.")
        try:
            fallback_concepts = self._fallback_extract_basic(text, filename, file_format)
            if fallback_concepts: # Check if fallback actually returned something
                logger.info(f"Successfully extracted {len(fallback_concepts)} concepts using fallback for {filename}.")
                return fallback_concepts
            else:
                logger.warning(f"Fallback extraction also yielded no concepts for {filename}.")
        except Exception as fallback_error:
            logger.error(f"Fallback extraction failed for {filename}: {str(fallback_error)}", exc_info=True)
        
        # If all else fails, generate default concepts
        logger.warning(f"All extraction methods failed for {filename}. Generating default concepts.")
        return self._generate_default_concepts(filename, file_format)

    def _extract_concepts_contextgem(self, text: str, filename: str, file_format: str) -> Optional[List[Dict]]:
        """Core concept extraction using ContextGem library. Returns None if ContextGem fails to produce valid/any concepts."""
        if not self.contextgem_llm:
            logger.error("ContextGem LLM is None in _extract_concepts_contextgem. Cannot proceed.")
            return None # Signal failure
            
        logger.info(f"Using ContextGem LLM ({self.contextgem_llm.model}) for concept extraction with ContextGem library.")
        
        cg_doc = ContextGemDocument(raw_text=text)
        cg_doc.add_concepts(self.default_concepts_for_cg)

        try:
            self.contextgem_llm.extract_concepts_from_document(cg_doc)

            extracted_concepts_list = []
            if hasattr(cg_doc, 'concepts') and cg_doc.concepts:
                for concept_obj in cg_doc.concepts:
                    if hasattr(concept_obj, 'name') and hasattr(concept_obj, 'extracted_items') and concept_obj.extracted_items:
                        for item in concept_obj.extracted_items:
                            item_value = getattr(item, 'value', str(item))
                            if item_value and str(item_value).strip():
                                extracted_concepts_list.append({
                                    "content": str(item_value),
                                    "concept_name": concept_obj.name,
                                    "type": "concept",
                                    "source_sentence": str(item_value)[:150] + "..." if len(str(item_value)) > 150 else str(item_value),
                                    "metadata": {
                                        "filename": filename,
                                        "format": file_format,
                                        "extraction_method": "contextgem_library"
                                    }
                                })
            
            if not extracted_concepts_list:
                logger.warning(f"ContextGem ran but extracted no valid concepts for {filename}.")
                return None # Signal that ContextGem produced no usable output

            logger.info(f"ContextGem successfully extracted {len(extracted_concepts_list)} concept items for {filename}.")
            return extracted_concepts_list

        except Exception as e:
            # This catches errors during the ContextGem call itself (e.g., API errors, internal ContextGem errors)
            logger.error(f"Exception during ContextGem's extract_concepts_from_document for {filename}: {str(e)}", exc_info=True)
            return None # Signal failure

    # _parse_domain_concepts_response, _parse_ollama_response, and _parse_fast_response
    # are removed as they were specific to the old direct LLM call approach.
    # ContextGem's output is handled directly in _extract_concepts_contextgem.

    def _generate_default_concepts(self, filename: str, file_format: str) -> List[Dict]:
        """Generate default concepts based on filename when extraction fails."""
        # Extract meaningful terms from filename
        name_without_ext = os.path.splitext(filename)[0]
        words = re.findall(r'\w+', name_without_ext)
        
        concepts = []
        
        # Use filename as a basic concept
        if len(words) > 1:
            concepts.append({
                "content": " ".join(words),
                "concept_name": "Document Topic",
                "type": "concept",
                "source_sentence": filename,
                "metadata": {
                    "filename": filename,
                    "format": file_format,
                    "extraction_method": "filename_fallback"
                }
            })
        
        # Add file format as a technical concept
        concepts.append({
            "content": f"Document in {file_format.upper()} format",
            "concept_name": "Technical Information",
            "type": "concept",
            "source_sentence": filename,
            "metadata": {
                "filename": filename,
                "format": file_format,
                "extraction_method": "filename_fallback"
            }
        })
        
        return concepts

    # _parse_fast_response removed.

    def _fallback_extract_basic(self, text: str, filename: str, file_format: str) -> List[Dict]:
        """Basic concept extraction fallback using direct LLM calls."""
        try:
            from litellm import completion
            
            logger.info(f"Using fallback extraction for {filename}")
            
            # Use a simpler, more direct prompt
            prompt = f"""You are an expert text analyst. Extract key concepts from the following text.
For each concept, provide its type, content, and a brief source/context from the text.
YOU MUST FOLLOW THIS FORMAT EXACTLY FOR EACH CONCEPT:
CONCEPT_TYPE: [Choose one: Key Definition, Main Idea, Important Fact, Technical Term, Key Person, Key Organization, Key Date]
CONTENT: [The extracted concept, definition, idea, fact, term, person, organization, or date]
SOURCE: [A very brief quote or reference from the text indicating where this was found]

Text:
{text[:15000]} # Limit text length for fallback prompt

Example:
CONCEPT_TYPE: Key Definition
CONTENT: Photosynthesis is the process by which green plants use sunlight, water, and carbon dioxide to create their own food.
SOURCE: "Photosynthesis is the process by which green plants use sunlight..."

Extract 5-10 concepts total. Ensure each concept block starts with 'CONCEPT_TYPE:'."""

            model_name_for_fallback = None
            api_base_for_fallback = None
            is_ollama_fallback = False

            if self.contextgem_llm and self.contextgem_llm.model:
                # Prefer the model already initialized for ContextGem
                full_model_name = self.contextgem_llm.model
                if full_model_name.startswith("ollama/"):
                    model_name_for_fallback = full_model_name
                    api_base_for_fallback = self.contextgem_llm.api_base or "http://localhost:11434"
                    is_ollama_fallback = True
                else: # Assuming OpenAI or other direct model names
                    model_name_for_fallback = full_model_name
                    api_base_for_fallback = self.contextgem_llm.api_base
            else:
                # Fallback if contextgem_llm wasn't initialized or has no model
                logger.warning("ContextGem LLM not available for fallback, attempting direct default.")
                # Attempt to use a default Ollama model if server is reachable
                try:
                    ollama_check_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
                    requests.get(f"{ollama_check_url}/api/tags", timeout=1) # Quick check
                    model_name_for_fallback = "ollama/mistral:7b" # A common default
                    api_base_for_fallback = ollama_check_url
                    is_ollama_fallback = True
                    logger.info(f"Using default Ollama model for fallback: {model_name_for_fallback}")
                except Exception:
                    if os.environ.get("OPENAI_API_KEY"):
                        model_name_for_fallback = "gpt-4o-mini" # OpenAI default
                        logger.info(f"Using default OpenAI model for fallback: {model_name_for_fallback}")
                    else:
                        logger.error("No usable model configuration for fallback extraction.")
                        return []
            
            if not model_name_for_fallback:
                logger.error("Could not determine a model for fallback extraction.")
                return []

            logger.info(f"Fallback extraction using model: {model_name_for_fallback}")
            response = completion(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                api_base=api_base,
                max_tokens=1500,
                temperature=0.3
            )
            
            response_text = response.choices[0].message.content
            return self._parse_fallback_response(response_text, filename, file_format)
            
        except Exception as e:
            logger.error(f"Fallback extraction failed: {str(e)}")
            return []

    def _parse_fallback_response(self, response_text: str, filename: str, file_format: str) -> List[Dict]:
        """Parse the fallback LLM response to extract concepts."""
        concepts = []
        
        # Split response into blocks and parse each one
        lines = response_text.split('\n')
        current_concept = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('CONCEPT_TYPE:'):
                # Save previous concept if it exists
                if current_concept.get('content'):
                    concepts.append(self._create_concept_dict(current_concept, filename, file_format))
                
                # Start new concept
                concept_type = line.replace('CONCEPT_TYPE:', '').strip()
                current_concept = {'concept_name': concept_type}
                
            elif line.startswith('CONTENT:'):
                content = line.replace('CONTENT:', '').strip()
                current_concept['content'] = content
                
            elif line.startswith('SOURCE:'):
                source = line.replace('SOURCE:', '').strip()
                current_concept['source_sentence'] = source
        
        # Add the last concept
        if current_concept.get('content'):
            concepts.append(self._create_concept_dict(current_concept, filename, file_format))
        
        return concepts

    def _create_concept_dict(self, concept_data: Dict, filename: str, file_format: str) -> Dict:
        """Create a standardized concept dictionary."""
        return {
            "content": concept_data.get('content', ''),
            "concept_name": concept_data.get('concept_name', 'Important Fact'),
            "type": "concept",
            "source_sentence": concept_data.get('source_sentence', ''),
            "metadata": {
                "filename": filename,
                "format": file_format,
                "extraction_method": "fallback"
            }
        }

    def validate_file(self, filename: str) -> bool:
        """Validate if the file has a supported format."""
        if not filename:
            return False
        file_extension = os.path.splitext(filename)[1][1:].lower()
        return file_extension in self.supported_formats

    def process(self, content: bytes, source: str, custom_filename: str = None,
                stop_signal_check: Optional[callable] = None,
                chunk_processed_callback: Optional[callable] = None) -> Dict:
        """Main processing method - simplified and optimized, with stop signal and progress callback."""
        logger.info(f"Processing document: {source}")
        
        try:
            # Determine file type and extract text
            if os.path.exists(source):
                filename = os.path.basename(source)
                if not self.validate_file(filename):
                    raise ValueError(f"Unsupported file format. Supported: {', '.join(self.supported_formats)}")
                
                file_format = os.path.splitext(filename)[1][1:].lower()
                with open(source, 'rb') as f:
                    file_content = f.read()
                text = self.convert_to_text(file_content, file_format)
            else:
                # Direct input processing
                # First, determine the file format
                file_format = "txt"  # default
                
                # Check if it's a PDF by looking at source filename or content
                if source.lower().endswith('.pdf') or (isinstance(content, bytes) and content.startswith(b'%PDF')):
                    file_format = "pdf"
                elif source.lower().endswith('.docx'):
                    file_format = "docx"
                elif source.lower().endswith('.html'):
                    file_format = "html"
                
                # Extract text based on detected format
                if isinstance(content, bytes):
                    # Always convert bytes to text using the appropriate method
                    text = self.convert_to_text(content, file_format)
                else:
                    # String content
                    text = str(content)
                
                # Generate filename if not provided
                if custom_filename and custom_filename.strip():
                    filename = custom_filename.strip()
                    # Ensure filename has correct extension
                    if not filename.lower().endswith(f'.{file_format}'):
                        filename += f'.{file_format}'
                else:
                    # Generate filename from first few words of extracted text
                    words = text.split()[:5]  # Take first 5 words
                    if words:
                        filename = "_".join(word.strip('.,!?;:"()[]{}') for word in words)
                        filename = re.sub(r'[^\w\s-]', '', filename)  # Remove special chars
                        filename = re.sub(r'\s+', '_', filename)  # Replace spaces with underscores
                        filename = filename[:50] + f".{file_format}"  # Limit length and add extension
                    else:
                        filename = f"document.{file_format}"
            
            if not text.strip():
                raise ValueError("Document appears to be empty or could not be processed.")
            
            # Log processing stats
            token_count = self.count_tokens(text) # Full document token count
            word_count = len(text.split()) # Full document word count
            logger.info(f"Document stats - Full Tokens: {token_count}, Full Words: {word_count}")

            chunks = self.smart_chunk_text(text)
            logger.info(f"Document split into {len(chunks)} chunk(s) for concept extraction.")
            total_chunks = len(chunks)
            
            all_extracted_concepts = []
            for i, chunk in enumerate(chunks):
                if stop_signal_check and stop_signal_check():
                    logger.info("Document processing stopped by user signal during chunk processing.")
                    break # Exit the chunk processing loop

                logger.info(f"Processing chunk {i+1}/{total_chunks} for concepts...")
                chunk_concepts = self.extract_concepts(chunk, filename, file_format) # Pass chunk here
                if chunk_concepts:
                    all_extracted_concepts.extend(chunk_concepts)
                
                processed_count = i + 1
                logger.info(f"Chunk {processed_count}/{total_chunks} yielded {len(chunk_concepts if chunk_concepts else [])} concept items.")
                if chunk_processed_callback:
                    chunk_processed_callback(processed_count, total_chunks)
            
            # Remove duplicate concepts that might arise from chunking (simple check based on content and name)
            unique_concepts_dict = {}
            for concept in all_extracted_concepts:
                key = (concept.get("concept_name"), concept.get("content"))
                if key not in unique_concepts_dict:
                    unique_concepts_dict[key] = concept
            
            concepts = list(unique_concepts_dict.values())
            logger.info(f"Total unique concept items extracted from all chunks: {len(concepts)}")
            
            # Create metadata compatible with Streamlit app
            paragraphs = text.split('\n\n') # Based on full text
            metadata = {
                'total_words': word_count,
                'total_tokens': token_count,
                'total_paragraphs': len([p for p in paragraphs if p.strip()]),
                'concept_count': len(concepts),
                'filename': filename,
                'format': file_format,
                'chapters': [],  # Simple implementation - could be enhanced
                'sections': []   # Simple implementation - could be enhanced
            }
            
            return {
                "content": text,
                "filename": filename,
                "format": file_format,
                "concepts": concepts,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            raise
            
    def convert_to_text(self, content: bytes, file_format: str) -> str:
        """Convert various file formats to plain text - optimized."""
        if not content:
            return ""
            
        if not isinstance(content, bytes):
            if isinstance(content, str):
                return content
            try:
                content = bytes(content)
            except Exception as e:
                logger.error(f"Failed to convert content to bytes: {str(e)}")
                return ""
        
        try:
            if file_format == 'pdf':
                return self._extract_text_from_pdf(content)
            elif file_format == 'docx':
                return self._extract_text_from_docx(content)
            elif file_format == 'html':
                return self._extract_text_from_html(content.decode('utf-8', errors='ignore'))
            elif file_format == 'txt':
                return content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file format: {file_format}")
                
        except Exception as e:
            logger.error(f"Error converting {file_format} to text: {str(e)}")
            raise
            
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        text_parts = []
        
        # Check if content might be binary PDF data
        if content.startswith(b'%PDF'):
            # Process as binary PDF
            with io.BytesIO(content) as pdf_file:
                try:
                    with pdfplumber.open(pdf_file) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                            
                            # Log progress for large documents
                            if (page_num + 1) % 10 == 0:
                                logger.debug(f"Processed {page_num + 1} pages")
                                
                except Exception as e:
                    logger.error(f"PDF extraction error with pdfplumber: {str(e)}")
                    # Try fallback method if pdfplumber fails
                    try:
                        # If pdfplumber fails, try to decode as UTF-8 text
                        # This would happen if the PDF contains mostly text that's already UTF-8 encoded
                        decoded_text = content.decode('utf-8', errors='ignore')
                        # Only return the decoded text if it looks like valid text (not binary garbage)
                        if '%PDF' in decoded_text and not re.search(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]{10,}', decoded_text):
                            return decoded_text
                    except Exception as fallback_error:
                        logger.error(f"Fallback PDF text extraction also failed: {str(fallback_error)}")
                    
                    # If all else fails, raise the original error
                    raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        else:
            # If it doesn't start with %PDF, try to decode as text directly
            try:
                decoded_text = content.decode('utf-8', errors='ignore')
                return decoded_text
            except Exception as e:
                logger.error(f"Failed to decode content as UTF-8: {str(e)}")
                # Continue with pdfplumber as a fallback
                with io.BytesIO(content) as pdf_file:
                    try:
                        with pdfplumber.open(pdf_file) as pdf:
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text_parts.append(page_text)
                    except Exception as pdf_error:
                        logger.error(f"PDF extraction fallback error: {str(pdf_error)}")
                        raise ValueError(f"Failed to extract text from content: {str(e)}")
        
        # If we extracted text using pdfplumber, join the parts
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            # If no text was extracted but no errors were raised, return empty string
            logger.warning("No text extracted from PDF content")
            return ""
        
    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        text_parts = []
        
        with io.BytesIO(content) as docx_file:
            try:
                doc = Document(docx_file)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
            except Exception as e:
                logger.error(f"DOCX extraction error: {str(e)}")
                raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return "\n".join(text_parts)
        
    def _extract_text_from_html(self, content: str) -> str:
        """Extract text from HTML content."""
        try:
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for element in soup(["script", "style", "meta", "link"]):
                element.extract()
            
            # Get text and clean it up
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return '\n'.join(chunk for chunk in chunks if chunk)
            
        except Exception as e:
            logger.error(f"HTML extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from HTML: {str(e)}")

    def get_content_summary(self, processed_content: Dict) -> str:
        """Generate a content summary for display purposes."""
        try:
            content = processed_content.get('content', '')
            metadata = processed_content.get('metadata', {})
            concepts = processed_content.get('concepts', [])
            
            # Basic statistics
            word_count = metadata.get('total_words', len(content.split()))
            concept_count = len(concepts)
            
            # Generate summary
            lines = [
                f"Document: {processed_content.get('filename', 'Unknown')}",
                f"Format: {processed_content.get('format', 'Unknown').upper()}",
                f"Total Words: {word_count:,}",
                f"Concepts Extracted: {concept_count}",
                ""
            ]
            
            # Add concept breakdown if available
            if concepts:
                concept_types = {}
                for concept in concepts:
                    concept_type = concept.get('concept_name', 'Unknown')
                    concept_types[concept_type] = concept_types.get(concept_type, 0) + 1
                
                lines.append("Concept Breakdown:")
                for concept_type, count in concept_types.items():
                    lines.append(f"  • {concept_type}: {count}")
                lines.append("")
            
            # Add content preview
            preview_length = min(200, len(content))
            lines.append(f"Content Preview ({preview_length} chars):")
            lines.append(f"{content[:preview_length]}{'...' if len(content) > preview_length else ''}")
            
            return "\n".join(lines)
            
        except Exception as e:
            logger.error(f"Error generating content summary: {str(e)}")
            return f"Error generating summary: {str(e)}"

    def intelligent_segmentation(self, content: str) -> List[Dict]:
        """Provide intelligent segmentation of content into logical sections."""
        try:
            if not content or not isinstance(content, str):
                return []
            
            # Simple paragraph-based segmentation
            paragraphs = content.split('\n\n')
            segments = []
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    segments.append({
                        'id': i,
                        'type': 'paragraph',
                        'content': paragraph.strip(),
                        'word_count': len(paragraph.split()),
                        'start_position': content.find(paragraph),
                        'metadata': {
                            'segment_number': i + 1,
                            'total_segments': len([p for p in paragraphs if p.strip()])
                        }
                    })
            
            return segments
            
        except Exception as e:
            logger.error(f"Error in intelligent segmentation: {str(e)}")
            return []
